import streamlit as st
from docx import Document
from io import BytesIO
import pandas as pd

# Load the combined Excel file
file_path_combined = 'Combined_Questions.xlsx'
df_combined = pd.read_excel(file_path_combined)

# Load the new Excel file with additional data
file_path_latest = 'Elemen_dagang.xlsx'
df_latest = pd.read_excel(file_path_latest)

# Extract Question_Number as a list of integers
question_numbers = df_latest['Question_Number'].astype(int).tolist()
print("Question numbers:", question_numbers)

# Function to replace hashtags with user input
def replace_hashtags(doc_path, replacements):
    doc = Document(doc_path)
    for p in doc.paragraphs:
        for hashtag, replacement in replacements.items():
            if hashtag in p.text:
                p.text = p.text.replace(hashtag, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for hashtag, replacement in replacements.items():
                    if hashtag in cell.text:
                        cell.text = cell.text.replace(hashtag, replacement)
    return doc

# Function to convert document to BytesIO
def convert_to_bytes(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit form
st.title("Form Pengisian Data")

# Creating a dictionary to hold user inputs
replacements = {}
num_questions = len(df_combined)
num_pages = (num_questions // 7) + (1 if num_questions % 7 != 0 else 0)

# Allow user to select page
current_page = st.selectbox("Pilih halaman", list(range(1, num_pages + 1)), key="page_selector")

# Calculate start and end indices for the current page
start_index = (current_page - 1) * 7
end_index = min(start_index + 7, num_questions)

# Get subheader mappings from df_latest based on Question_Number
subheader_mappings = {int(row['Question_Number']): row['Capaian Pembelajaran'] for _, row in df_latest.iterrows()}

# Form for the current page
with st.form("data_form"):
    all_filled = True
    for i, (index, row) in enumerate(df_combined.iloc[start_index:end_index].iterrows(), start=1):
        actual_question_number = start_index + i + 2

        # Check if we need to insert a subheader
        if actual_question_number in question_numbers:
            subheader = subheader_mappings.get(actual_question_number)
            if subheader:
                st.subheader(subheader)
                print(f"Inserting subheader: {subheader} for question number {actual_question_number}")
            else:
                print(f"No subheader found for question number: {actual_question_number}")

        question = row['Pertanyaan']
        hashtag = row['Penanda'] if row['Source'] == "Data Diri" else row['Answer Yes']
        tipe = row['Tipe']

        print(f"Processing question {actual_question_number}")
        print(f"Question: {question}")
        print(f"Hashtag: {hashtag}")
        print(f"Type: {tipe}")
        print("---")

        response = None  # Initialize response variable

        if tipe == 'short description':
            response = st.text_input(question, key=hashtag)
            replacements[hashtag] = str(response)
        elif tipe == 'paragraph':
            response = st.text_area(question, key=hashtag)
            replacements[hashtag] = str(response)
        elif tipe == 'date':
            response = st.date_input(question, key=hashtag)
            replacements[hashtag] = str(response)
        elif tipe == 'radio':
            response = st.radio(question, ['Belum memilih', 'Sekolah', 'Iduka'], key=hashtag)

            if response == 'Ya':
                replacements['Answer Yes'] = 'Ya'
                replacements['Answer No'] = ''
            elif response == 'Tidak':
                replacements['Answer Yes'] = ''
                replacements['Answer No'] = 'Tidak'
            else:
                all_filled = False

        if response == '' or response is None:
            all_filled = False

        # Add a separator between questions
        st.markdown("---")

    # Place the submit button at the bottom
    submit_button = st.form_submit_button(label="Generate Document")

if submit_button:
    if all_filled:
        # Replace hashtags in the document
        doc_path = "Template_Document.docx"  # Ensure this file is in the same directory
        updated_doc = replace_hashtags(doc_path, replacements)

        # Convert document to BytesIO
        bytes_io = convert_to_bytes(updated_doc)

        # Provide download link
        st.download_button(label="Download file disini", data=bytes_io, file_name="Filled_Document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Harap selesaikan semua pertanyaan sebelum mengunduh proposal.")
